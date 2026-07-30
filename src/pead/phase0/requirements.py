"""Deterministic source-clause inventory for the PEAD specification."""

from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterator

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from pead.config.models import ConfigValidationError


EXPECTED_BODY_PARAGRAPHS = 578
EXPECTED_BODY_TABLES = 83


@dataclass(frozen=True)
class SectionPlan:
    phases: tuple[str, ...]
    files: tuple[str, ...]
    tests: tuple[str, ...]
    produced_artifact: str
    release_failure_condition: str
    affected_claims: tuple[str, ...]


DEFAULT_PLAN = SectionPlan(
    phases=("0", "13"),
    files=("WorkPlan.md", "configs/requirements/pead_v1_requirements.yaml"),
    tests=("scripts/audit_phase0.py", "future phase-specific audit"),
    produced_artifact="clause coverage and final evidence manifest",
    release_failure_condition="Unimplemented or unaudited normative clause blocks affected release.",
    affected_claims=("C1", "C2", "C3", "C4", "C5", "C6"),
)


SECTION_PLANS: dict[str, SectionPlan] = {
    "0": SectionPlan(("0",), ("CLAIMS.md", "configs/study/pead_main_v1.yaml"), ("phase-0 charter audit",), "frozen study charter", "Charter mismatch blocks Phase 0.", ("C1", "C2", "C3", "C4", "C5", "C6")),
    "1": SectionPlan(("0", "9", "13"), ("CLAIMS.md", "configs/study/pead_main_v1.yaml"), ("claim dependency and wording audit",), "claim ledger", "Claim mismatch or overclaim blocks release.", ("C1", "C2", "C3", "C4", "C5", "C6")),
    "2": SectionPlan(("0", "2", "3", "6", "9", "12"), ("configs/access/*.yaml", "src/pead/tracks/*", "src/pead/metrics/*"), ("equivalence, divergence, lower-bound, distance, access audits",), "formal and empirical PEAD reports", "Failure blocks C1-C3 and affected bank.", ("C1", "C2", "C3", "C6")),
    "3": SectionPlan(("0", "9", "13"), ("CLAIMS.md", "configs/requirements/pead_v1_requirements.yaml"), ("causal-rejection closure audit",), "closure report", "Missing control, audit, gate, or evidence blocks release.", ("C1", "C2", "C3", "C4", "C5", "C6")),
    "4": SectionPlan(("0", "1", "6", "11", "12"), ("src/pead/core/*", "src/pead/projections/*", "docs/blind_custody_protocol.md"), ("trust-boundary, access, trace-order, custody audits",), "trust-boundary report", "Boundary violation invalidates affected run.", ("C1", "C2", "C3", "C4", "C5", "C6")),
    "5": SectionPlan(("0", "1"), ("README.md", "pyproject.toml", "src/pead/*", "tests/*"), ("repository layout and import-boundary audit",), "repository structure report", "Missing required responsibility boundary blocks dependent phase.", ("C1", "C2", "C3", "C4", "C5", "C6")),
    "6": SectionPlan(("0", "1", "6"), ("configs/access/*.yaml", "src/pead/core/types.py", "src/pead/projections/*"), ("schema, canonicalization, mask, trace audits",), "data dictionary and access manifest", "State or trace contract failure blocks affected claims.", ("C1", "C2", "C3", "C4", "C5", "C6")),
    "7": SectionPlan(("3", "5", "9A"), ("src/pead/world/*", "configs/mechanisms/*", "configs/domains/*"), ("generator, intervention, nuisance, independence audits",), "generator and mechanism manifests", "Generator defect quarantines bank or invalidates study.", ("C1", "C2", "C3", "C4", "C5", "C6")),
    "8": SectionPlan(("2", "4", "9", "12"), ("src/pead/labels/*", "configs/policies/*"), ("dual-engine, Oracle, ambiguity, monotonicity audits",), "label and Oracle reports", "Label defect blocks case, bank, or release.", ("C1", "C2", "C3", "C6")),
    "9": SectionPlan(("3", "4", "9A", "12"), ("src/pead/tracks/*", "configs/tracks/*", "configs/allocations/*"), ("track allocation, integrity, distance, reversal, scope, evidence audits",), "track manifests and reports", "Track failure blocks its scientific claim.", ("C1", "C2", "C3", "C5", "C6")),
    "10": SectionPlan(("5", "9A"), ("src/pead/domains/*", "configs/domains/*"), ("domain schema, validity, non-triviality, custody audits",), "domain review and manifest", "Domain failure blocks broad or generalized claims.", ("C1", "C2", "C3", "C4", "C5", "C6")),
    "11": SectionPlan(("6", "7", "9"), ("src/pead/projections/*", "configs/access/*"), ("access, parity, budget, canary audits",), "access and fairness reports", "Forbidden access invalidates method results.", ("C2", "C3", "C4")),
    "12": SectionPlan(("0", "7", "10", "12"), ("configs/methods/method_inventory_v1.yaml", "src/pead/baselines/*", "src/pead/mavs/*"), ("method inventory, fidelity, budget, training audits",), "method and training manifests", "Method defect blocks affected comparison.", ("C2", "C3", "C4", "C5", "C6")),
    "13": SectionPlan(("0", "7", "9A", "10", "11", "12"), ("configs/holdouts/*", "docs/blind_custody_protocol.md", "src/pead/holdouts/*"), ("split, holdout, chronology, contamination, freeze, blind adjudication audits",), "holdout, freeze, and blind-run manifests", "Contamination or chronology failure invalidates blind claims.", ("C1", "C2", "C3", "C4", "C5", "C6")),
    "14": SectionPlan(("0", "9", "12", "13"), ("configs/metrics/*", "src/pead/metrics/*", "src/pead/reports/*"), ("metric fixtures, statistics, outcome-tier, claim audits",), "metric, statistical, and outcome reports", "Metric or inference defect blocks affected claim.", ("C1", "C2", "C3", "C4", "C5", "C6")),
    "15": SectionPlan(("0", "7", "8", "12"), ("configs/methods/method_inventory_v1.yaml", "src/pead/mavs/*"), ("ablation, scalar, equal-information, architecture audits",), "ablation and architecture report", "Failure limits or blocks C4.", ("C4", "C5")),
    "16": SectionPlan(("9A", "10", "12"), ("configs/allocations/*", "manifests/allocations/*", "results/*"), ("allocation, denominator, balance, volume, compute audits",), "allocation and execution manifests", "Allocation or budget failure blocks claim bank.", ("C1", "C2", "C3", "C4", "C5", "C6")),
    "17": SectionPlan(tuple(str(index) for index in range(14)), ("WorkPlan.md", "Path.md"), ("phase-close audit",), "phase evidence ledger", "Incomplete phase remains open.", ("C1", "C2", "C3", "C4", "C5", "C6")),
    "18": SectionPlan(("0", "1", "2", "3", "4", "5", "6", "7", "8", "9", "9A", "10", "11", "12", "13"), ("src/pead/audits/*", "scripts/audit_*.py", "tests/*"), ("unit, property, metamorphic, integration, stress, master audits",), "machine and human audit reports", "Release-blocking audit failure blocks release.", ("C1", "C2", "C3", "C4", "C5", "C6")),
    "19": SectionPlan(("0", "13"), ("REPRODUCE.md", "requirements.lock", "manifests/*"), ("clean reproduction and artifact traceability audits",), "public evidence and reproduction package", "Reproduction or traceability failure blocks release.", ("C1", "C2", "C3", "C4", "C5", "C6")),
    "20": SectionPlan(("0", "9", "13"), ("CLAIMS.md", "CLAIM_ELIGIBILITY.md", "src/pead/reports/*"), ("claim predicate, forbidden wording, paper-boundary audits",), "claim eligibility map", "Overclaim or paper-boundary failure blocks release.", ("C1", "C2", "C3", "C4", "C5", "C6")),
    "21": SectionPlan(("0", "9", "12", "13"), ("configs/study/pead_main_v1.yaml", "src/pead/audits/*"), ("stop-condition fixtures and master audit",), "stop-condition report", "Triggered mandatory stop invalidates affected execution.", ("C1", "C2", "C3", "C4", "C5", "C6")),
    "22": SectionPlan(("0", "9", "13"), ("CLAIMS.md", "configs/requirements/pead_v1_requirements.yaml"), ("causal-rejection closure audit",), "closure and residual-uncertainty report", "Unclosed preventable concern blocks release.", ("C1", "C2", "C3", "C4", "C5", "C6")),
    "23": SectionPlan(("0", "13"), ("Path.md", "configs/requirements/pead_v1_requirements.yaml"), ("final completeness audit",), "final audit and revision record", "Missing or weakened requirement blocks release.", ("C1", "C2", "C3", "C4", "C5", "C6")),
    "A": SectionPlan(("0", "1", "4", "6", "7"), ("configs/study/failure_card_schema_v1.yaml", "src/pead/core/types.py"), ("record schema tests",), "data dictionary", "Missing field blocks dependent artifact.", ("C1", "C2", "C3", "C4", "C5", "C6")),
    "B": SectionPlan(("1", "3", "9", "12", "13"), ("src/pead/core/*", "src/pead/metrics/*", "src/pead/reports/*"), ("algorithm contract tests",), "algorithm audit report", "Algorithm contract failure blocks affected result.", ("C1", "C2", "C3", "C4")),
    "C": SectionPlan(("0", "1", "2", "3", "4", "5", "6", "7", "8", "9", "10", "11", "12", "13"), ("tests/*",), ("continuous integration inventory audit",), "test reports", "Missing required test blocks affected phase.", ("C1", "C2", "C3", "C4", "C5", "C6")),
    "D": SectionPlan(("0", "9A", "11", "12", "13"), ("manifests/*",), ("manifest completeness and hash audits",), "final artifact manifest", "Missing identity blocks release.", ("C1", "C2", "C3", "C4", "C5", "C6")),
}


def normalize_text(value: str) -> str:
    return " ".join(value.split())


def sha256_text(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def iter_body_blocks(document: DocumentObject) -> Iterator[Paragraph | Table]:
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def section_key(heading: str) -> str:
    appendix = re.match(r"Appendix\s+([A-D])\b", heading)
    if appendix:
        return appendix.group(1)
    numbered = re.match(r"(\d+)(?:\.|\s)", heading)
    if numbered:
        return numbered.group(1)
    return "0"


def classify_clause(text: str, style_name: str, heading: str, is_table: bool) -> str:
    lower = f"{heading} {text}".lower()
    if is_table:
        return "table_row"
    if style_name.startswith("List"):
        if "stop condition" in heading.lower():
            return "stop_condition"
        if "completion gate" in lower or "release gate" in lower:
            return "gate"
        return "normative_bullet"
    if "metric" in lower:
        return "metric_clause"
    if "audit" in lower:
        return "audit_clause"
    if re.search(r"\b(must|shall|required|cannot|may not|only if|blocks?|invalidates?)\b", lower):
        return "normative_paragraph"
    return "supporting_specification_clause"


def build_registry(source_docx: Path) -> dict[str, Any]:
    source_docx = source_docx.resolve()
    if not source_docx.is_file():
        raise ConfigValidationError(f"Source DOCX does not exist: {source_docx}")
    document = Document(source_docx)
    source_hash = hashlib.sha256(source_docx.read_bytes()).hexdigest()
    if len(document.paragraphs) != EXPECTED_BODY_PARAGRAPHS:
        raise ConfigValidationError(
            f"Unexpected paragraph count: {len(document.paragraphs)} "
            f"(expected {EXPECTED_BODY_PARAGRAPHS})"
        )
    if len(document.tables) != EXPECTED_BODY_TABLES:
        raise ConfigValidationError(
            f"Unexpected table count: {len(document.tables)} "
            f"(expected {EXPECTED_BODY_TABLES})"
        )
    paragraph_index = -1
    table_index = -1
    heading = "Document root"
    requirements: list[dict[str, Any]] = []
    heading_records: list[dict[str, str]] = []
    for block in iter_body_blocks(document):
        if isinstance(block, Paragraph):
            paragraph_index += 1
            text = normalize_text(block.text)
            if not text:
                continue
            style_name = block.style.name if block.style is not None else ""
            if style_name.startswith("Heading"):
                heading = text
                heading_records.append(
                    {
                        "source_locator": f"paragraph:{paragraph_index}",
                        "heading": heading,
                        "sha256": sha256_text(heading),
                    }
                )
                continue
            key = section_key(heading)
            plan = SECTION_PLANS.get(key, DEFAULT_PLAN)
            requirement_id = f"DOCX-P{paragraph_index:04d}"
            requirements.append(
                {
                    "requirement_id": requirement_id,
                    "source_locator": f"paragraph:{paragraph_index};heading:{heading}",
                    "exact_source_clause": text,
                    "source_clause_sha256": sha256_text(text),
                    "normative_class": classify_clause(text, style_name, heading, False),
                    "phases": list(plan.phases),
                    "files": list(plan.files),
                    "tests": list(plan.tests),
                    "produced_artifact": plan.produced_artifact,
                    "release_failure_condition": plan.release_failure_condition,
                    "affected_claims": list(plan.affected_claims),
                }
            )
        else:
            table_index += 1
            key = section_key(heading)
            plan = SECTION_PLANS.get(key, DEFAULT_PLAN)
            for row_index, row in enumerate(block.rows):
                cell_text = [normalize_text(cell.text) for cell in row.cells]
                text = " | ".join(cell_text)
                if not text.replace("|", "").strip():
                    continue
                requirement_id = f"DOCX-T{table_index:03d}-R{row_index:03d}"
                requirements.append(
                    {
                        "requirement_id": requirement_id,
                        "source_locator": (
                            f"table:{table_index};row:{row_index};heading:{heading}"
                        ),
                        "exact_source_clause": text,
                        "source_clause_sha256": sha256_text(text),
                        "normative_class": classify_clause(text, "", heading, True),
                        "phases": list(plan.phases),
                        "files": list(plan.files),
                        "tests": list(plan.tests),
                        "produced_artifact": plan.produced_artifact,
                        "release_failure_condition": plan.release_failure_condition,
                        "affected_claims": list(plan.affected_claims),
                    }
                )
    requirement_ids = [item["requirement_id"] for item in requirements]
    if len(requirement_ids) != len(set(requirement_ids)):
        raise ConfigValidationError("Generated requirement IDs are not unique")
    clauses_digest = hashlib.sha256(
        "\n".join(
            f"{item['requirement_id']}:{item['source_clause_sha256']}"
            for item in requirements
        ).encode("utf-8")
    ).hexdigest()
    headings_digest = hashlib.sha256(
        "\n".join(
            f"{item['source_locator']}:{item['sha256']}" for item in heading_records
        ).encode("utf-8")
    ).hexdigest()
    return {
        "schema_version": "1.0",
        "registry_id": "PEAD-REQUIREMENTS-v1",
        "source_document": {
            "filename": source_docx.name,
            "sha256": source_hash,
            "body_paragraph_count": len(document.paragraphs),
            "body_table_count": len(document.tables),
        },
        "extraction_contract": {
            "included": [
                "every nonempty non-heading body paragraph",
                "every nonempty table row including headers and callouts",
            ],
            "excluded": ["empty paragraphs", "section headings recorded separately"],
            "normalization": "Unicode text with all whitespace runs collapsed to one ASCII space",
            "stable_id_rules": [
                "DOCX-P#### uses the zero-based top-level paragraph index",
                "DOCX-T###-R### uses zero-based top-level table and row indices",
            ],
        },
        "heading_count": len(heading_records),
        "heading_inventory_sha256": headings_digest,
        "included_clause_count": len(requirements),
        "clause_inventory_sha256": clauses_digest,
        "required_entry_fields": [
            "requirement_id",
            "source_locator",
            "exact_source_clause",
            "source_clause_sha256",
            "normative_class",
            "phases",
            "files",
            "tests",
            "produced_artifact",
            "release_failure_condition",
            "affected_claims",
        ],
        "requirements": requirements,
    }
